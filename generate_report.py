from docx import Document
from datetime import datetime
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.shared import Inches, RGBColor
import matplotlib
matplotlib.use("Agg")  
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from zapv2 import ZAPv2
import matplotlib.pyplot as plt
from copy import deepcopy
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from google import genai
import time, json, os
from dotenv import load_dotenv
from extensions import db
from models import Vulnerabilidades_totales
load_dotenv()

GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
ZAP_API_KEY = os.getenv("ZAP_API_KEY")
ZAP_URL = os.getenv("ZAP_URL")
zap_url = ZAP_URL

doc = Document(r"C:\Users\gizquierdog\Documents\custom_report\custom_report.docx")

def connect_zap():
    try:
        zap = ZAPv2(apikey=ZAP_API_KEY, proxies={'http': zap_url, 'https': zap_url})
        print(f"Conectado a ZAP, versión: {zap.core.version}")
        return zap
    except Exception as e:
        print(f"Error al conectar con ZAP: {e}")
        return None




def remplazar_texto(doc, remplazos):
    for paragraph in doc.paragraphs:
        for marcador, valor in remplazos.items():
            if marcador in paragraph.text:
                paragraph.text = paragraph.text.replace(marcador, valor)


def remplazar_encabezado(doc, remplazos):
    for section in doc.sections:
        header = section.header
        for table in header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for marcador, valor in remplazos.items():
                        if marcador in cell.text:
                            cell.text = cell.text.replace(marcador, valor)
                            for paragraph in cell.paragraphs:
                                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                for run in paragraph.runs:
                                    run.font.color.rgb = RGBColor(255, 255, 255) 

def modificar_primer_tabla(doc, remplazos):
    tabla = doc.tables[0]  
    for row in tabla.rows:
        for cell in row.cells:
            for marcador, valor in remplazos.items():
                if marcador in cell.text:
                    cell.text = cell.text.replace(marcador, valor)
                    for parrafo in cell.paragraphs:
                        parrafo.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

def grafica_barras(vul_altas, vul_medias, vul_bajas, vul_informativas):
    
    labels = ["Alta", "Media", "Baja", "Informativa"]
    values = [vul_altas, vul_medias, vul_bajas, vul_informativas]
    colors = ['#E57373', '#FFB74D', '#FFF176', '#64B5F6']  # Colores más suaves

    # Crear la figura
    fig, ax = plt.subplots(figsize=(20, 1.2))  # Gráfica más alargada y menos ancha

    # Dibujar las barras horizontales segmentadas
    ax.barh([0], [values[0]], color=colors[0], height=0.4, left=0)  # Primera barra
    ax.barh([0], [values[1]], color=colors[1], height=0.4, left=values[0])  # Segunda barra
    ax.barh([0], [values[2]], color=colors[2], height=0.4, left=sum(values[:2]))  # Tercera barra
    ax.barh([0], [values[3]], color=colors[3], height=0.4, left=sum(values[:3]))  # Cuarta barra

    # Personalización del diseño
    ax.set_xlim(0, sum(values))  # Ajustar el límite del eje X
    ax.set_yticks([])  # Ocultar el eje Y
    ax.set_xticks([])  # Ocultar el eje X

    for i, v in enumerate(values):
        x_pos = sum(values[:i]) + v / 2  # Posición centrada en cada segmento
        ax.text(x_pos, 0, str(v), color='black', va='center', ha='center', fontsize=15)  # Texto dentro de las barras

    # Eliminar bordes y ajustar diseño limpio
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['left'].set_visible(False)
    ax.spines['bottom'].set_visible(False)

    # Añadir título completo
    plt.title("Vulnerabilidades por Criticidad", fontsize=30, loc='center')

    # Añadir leyenda debajo de la gráfica
    legend_labels = ["Alta", "Media", "Baja", "Informativa"]
    legend_colors = ['#E57373', '#FFB74D', '#FFF176', '#64B5F6']
    for i, label in enumerate(legend_labels):
        plt.plot([], [], color=legend_colors[i], label=label)

    plt.legend(loc='lower center', bbox_to_anchor=(0.5, -1.25), ncol=len(legend_labels), fontsize=30)

    image_path = (r"C:\Users\gizquierdog\Documents\custom_report\grafica_vulnerabilidades.png")

    # Guardar la gráfica como imagen
    plt.tight_layout()
    plt.savefig(image_path,bbox_inches='tight')
    plt.close()

    return image_path



def insertar_imagen_en_celda(doc, image_path, tabla_index, fila, columna, width=Inches(7)):
    try:
        # Comprobar que existe la tabla especificada
        if tabla_index >= len(doc.tables):
            print(f" Error: No existe la tabla {tabla_index}. Hay {len(doc.tables)} tablas en el documento.")
            return

        tabla = doc.tables[tabla_index]

        # Comprobar que existen la fila y columna especificadas
        if fila >= len(tabla.rows) or columna >= len(tabla.columns):
            print(f"Error: La celda [{fila}, {columna}] está fuera del rango de la tabla.")
            return

        celda = tabla.cell(fila, columna)

        # Limpiar contenido previo de la celda (opcional)
        celda.text = ""

        # Insertar imagen en la celda
        paragraph = celda.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(image_path, width=width)

        print(f"Imagen insertada correctamente en la tabla {tabla_index}, fila {fila}, columna {columna}")

    except Exception as e:
        print(f"Error inesperado al insertar la imagen: {e}")
##########################################################################################################################################################33
     
def agregar_alerta_tabla_6(doc, datos_alerta):
    tabla = doc.tables[6]
    # Agregar una nueva fila al final de la tabla
    nueva_fila = tabla.add_row()
    
    color_hex = "#FFFFFF"  
    # Completar las celdas de la nueva fila con los datos proporcionados
    celdas = nueva_fila.cells
    celdas[0].text = str(datos_alerta[0])  # Vulnerabilidad
    celdas[1].text = str(datos_alerta[1])  # Número de alertas
    celdas[2].text = str(datos_alerta[2])  # Categoría OWASP 
    # Criticidad (texto negro y fondo con color)
    celdas[3].text = str(datos_alerta[3])
    celdas[3].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)  # Texto negro 
    celdas[4].text = str(datos_alerta[4])
    celdas[4].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)

    if datos_alerta[3] == "Alta" or datos_alerta[3]== "High":
        color_hex = "#D32F2F"  # Rojo suave
    elif datos_alerta[3] == "Media" or datos_alerta[3] == "Medium":
        color_hex = "#F57C00"  # Naranja suave
    elif datos_alerta[3] == "Baja" or datos_alerta[3] == "Low":
        color_hex = "#FBC02D"  # Amarillo suave
    elif datos_alerta[3] == "Informativa" or datos_alerta[3] == "Informational":
        color_hex = "#1976D2"  # Azul suave

    # Aplicar color de fondo a la celda
    celdas[3]._element.get_or_add_tcPr().append(parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{(color_hex or "#FFFFFF").replace("#", "")}"/>'
    ))

    for cell in nueva_fila.cells:
        tc_pr = cell._element.get_or_add_tcPr()
        tc_pr.append(parse_xml(
            f'''<w:tcBorders {nsdecls('w')}>
                    <w:top w:val="nil"/>
                    <w:left w:val="nil"/>
                    <w:bottom w:val="nil"/>
                    <w:right w:val="nil"/>
                </w:tcBorders>'''
        ))
def consulta_gemini(alert_name,alert_desc,alert_cwe):
    prompt_text = f"""
    Eres un asistente especializado en ciberseguridad. Se te proporcionará el nombre de una alerta y su descripción.

    Tu tarea:
    Generar una nueva descripción clara y concisa de la alerta. (45-60 palabras)
    Identificar los posibles riesgos asociados a esta vulnerabilidad.(45-60 palabras)
    Proporcionar una solución para mitigar el problema.(45-60 palabras)
    Determinar el OWASP Top 10 correspondiente, únicamente basándote en el CWE asociado, Solo quiero el A01, A02...
    Formato de salida (JSON):
    json
    "detalles": "",
    "riesgo": "",
    "solucion": "",
    "owasp": ""
    alerta: {alert_name}
    descripcion: {alert_desc} 
    CWE: {alert_cwe}"""

    client = genai.Client(api_key=GEMINI_API_KEY)
    try:
        response = client.models.generate_content(
            model="gemini-2.0-flash", contents=prompt_text,
            config={
                'response_mime_type': 'application/json',
            },
        )
        return json.loads(response.text)
    except Exception as e:
        print(f"Error al obtener datos de la alerta: {e}")
        return {"descripcion": "", "riesgo": "", "solucion": "", "owasp": ""}
    
def agragar_datos_owasp_vulneravilidades_totales(owaps_top_10):
    vuln_totales = db.session.query(Vulnerabilidades_totales).first()
    if owaps_top_10 == 'A01':
        vuln_totales.a01 += 1
    elif owaps_top_10 == 'A02':
        vuln_totales.a02 += 1
    elif owaps_top_10 == 'A03':
        vuln_totales.a03 += 1
    elif owaps_top_10 == 'A04':
        vuln_totales.a04 += 1
    elif owaps_top_10 == 'A05': 
        vuln_totales.a05 += 1
    elif owaps_top_10 == 'A06':
        vuln_totales.a06 += 1
    elif owaps_top_10 == 'A07':
        vuln_totales.a07 += 1
    elif owaps_top_10 == 'A08':
        vuln_totales.a08 += 1
    elif owaps_top_10 == 'A09':
        vuln_totales.a09 += 1
    elif owaps_top_10 == 'A10':
        vuln_totales.a10 += 1
    db.session.commit()

def get_alertas(url):
    zap = connect_zap()
    alerts = zap.alert.alerts(url)
    risk_order = {"High": 1, "Medium": 2, "Low": 3, "Informational": 4}
    risk_translation = {
    "High": "Alto",
    "Medium": "Medio",
    "Low": "Bajo",
    "Informational": "Informativo"
    }
    alerts_sorted = sorted(alerts, key=lambda x: risk_order.get(x.get('risk', 'Informational'), 4))
    alertas_set = set()
    alertas_high_set = set()
    alertas_medium_set = set()
    alertas_low_set = set()
    alertas_informational_set = set()
    cont = 1
    
    alertas_info =[]
    for alert in alerts_sorted:
        alert_name = alert.get('name')
        if alert.get('risk') == 'High':
            alertas_high_set.add(alert_name)
        elif alert.get('risk') == 'Medium':
            alertas_medium_set.add(alert_name)
        elif alert.get('risk') == 'Low':
            alertas_low_set.add(alert_name)    
        elif alert.get('risk') == 'Informational':
            alertas_informational_set.add(alert_name)
        
        if alert_name in alertas_set:
            continue
        
        alertas_set.add(alert_name)
        alert_risk = alert.get('risk')
        alert_risk_spanish = risk_translation.get(alert_risk, alert_risk)
        alertas_filtradas = [alerta for alerta in alerts if alerta['alert'] == alert_name]
        alert_count = len(alertas_filtradas)
        alert_desc = alert.get('desc')
        alert_cwe = alert.get('cweid')
        alert_references = alert.get('reference')
        datos = consulta_gemini(alert_name,alert_desc,alert_cwe)
        time.sleep(0.5)
        alertas_info.append({
            'numero': f"{cont:02d}",
            'alert_name': alert_name,
            'risk': alert_risk_spanish,
            'owasp': datos['owasp'], 
            'cwe': alert_cwe,
            'url': url,
            'detalles': datos['detalles'],
            'riesgo': datos['riesgo'],
            'solucion': datos['solucion'],
            'referencias':alert_references
        })
        datos_alerta = [f"[VUL 0{cont}] {alert_name}", alert_count, datos["owasp"], alert_risk_spanish, "Detectada"]
        agragar_datos_owasp_vulneravilidades_totales(datos_alerta[2])
        cont += 1
        agregar_alerta_tabla_6(doc, datos_alerta)
    agregar_tablas_vulnerabilidades(doc,len(alertas_set))
    tabla_index = 8
    print(alertas_high_set,alertas_medium_set,alertas_low_set,alertas_informational_set)
    for i, alert_info in enumerate(alertas_info):
        rellenar_tabla_vulnerabilidades(doc,tabla_index+i,alert_info)
    
    return  alertas_set, alertas_high_set, alertas_medium_set, alertas_low_set, alertas_informational_set


def rellenar_tabla_vulnerabilidades(doc,cont,alert_info):
    tabla = doc.tables[cont]
    for row in tabla.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for key, value in alert_info.items():
                    key_formatted = "{" + key + "}"
                    if key_formatted in paragraph.text:
                        # Reemplazar manteniendo el formato
                        for run in paragraph.runs:
                            if key_formatted in run.text:
                                run.text = run.text.replace(key_formatted, value)
def agregar_tablas_vulnerabilidades(doc,n):
    """
    Copia una tabla existente y la agrega en una nueva página inmediatamente después de la tabla original.
    """
    tabla_original = doc.tables[8]
    for i in range(n-1):
        salto_pagina = OxmlElement("w:p")  
        run = OxmlElement("w:r")  
        break_tag = OxmlElement("w:br") 
        
        break_tag.set(qn("w:type"), "page") 

        run.append(break_tag)  
        salto_pagina.append(run)  
        
        nueva_tabla = deepcopy(tabla_original._element)

        tabla_original._element.addnext(salto_pagina)  
        salto_pagina.addnext(nueva_tabla) 
        
    doc.save(r"C:\Users\gizquierdog\Documents\custom_report\custom_report_modificado.docx")
def contexto_resumen_ejecutivo(url, alertas_set, target_url):
    datos_json = alertas_set
    time.sleep(1)
    prompt_text = f"""
        Eres un experto en ciberseguridad y redacción de informes ejecutivos. Tu tarea es generar un resumen ejecutivo claro y accesible basado en un conjunto de vulnerabilidades de seguridad detectadas en una auditoría.

        Instrucciones:

        Redacta un informe ejecutivo de aproximadamente 210 palabras.
        No incluyas títulos ni encabezados.
        El informe debe destacar las vulnerabilidades más importantes, las recomendaciones clave y el estado general de seguridad.
        Usa un lenguaje accesible para que cualquier persona, sin conocimientos técnicos, pueda comprenderlo fácilmente.
        No te enfoques en clasificaciones como "baja" o "informativa"; prioriza el impacto real y el riesgo asociado a cada alerta.
        No exageres la gravedad de las vulnerabilidades; describe cada una con la importancia que merece.
        La redacción debe ser profesional, objetiva y concisa.
        Datos de entrada:
        Aquí tienes el conjunto de vulnerabilidades detectadas en formato JSON:
        {datos_json}
        URL asociada al análisis: {url}

        Ejemplo del estilo esperado:

        En la auditoría de seguridad realizada en [nombre de la plataforma], se identificaron vulnerabilidades que pueden comprometer la confidencialidad y disponibilidad del sistema. Destaca una vulnerabilidad crítica que permite la ejecución de código malicioso en la aplicación, lo que podría ser explotado por atacantes. Además, se detectaron deficiencias en los mecanismos de autenticación, como la ausencia de autenticación multifactor, incrementando el riesgo de accesos no autorizados. También se encontraron configuraciones inseguras en la gestión de cookies y cabeceras HTTP, lo que podría exponer información sensible. Aunque existen medidas de protección implementadas, es fundamental fortalecer los controles de acceso y mitigar las vulnerabilidades más críticas para mejorar la seguridad general."""
    client = genai.Client(api_key=GEMINI_API_KEY)
    try:
        response = client.models.generate_content(
            model="gemini-2.0-flash", contents=prompt_text,
        )

        remplazos = {
            "{url}": target_url,
            "{resumen_ejecutivo}": response.text
        }
        for para in doc.paragraphs:
            for clave, valor in remplazos.items():
                if clave in para.text:
                    para.text = para.text.replace(clave, valor)

    except Exception as e:
        print(f"Error al obtener datos de la alerta: {e}")
        return {"descripcion": "", "riesgo": "", "solucion": "", "owasp": ""}


def generar_reporte_custom(target_url):
    remplazos = {
        "{nombre-url}": target_url,
        "{date}": datetime.now().strftime('%d/%m/%Y'),
    }
    remplazar_texto(doc, remplazos)
    remplazar_encabezado(doc, remplazos)
    modificar_primer_tabla(doc, remplazos)
    alertas_set, high_set, medium_set, low_set, informational_set = get_alertas(target_url) 
    contexto_resumen_ejecutivo(target_url, alertas_set, target_url)
    imagen_path = grafica_barras(len(high_set), len(medium_set), len(low_set), len(informational_set))
    insertar_imagen_en_celda(doc, imagen_path, tabla_index=5, fila=1, columna=0)
    doc.save(r"C:\Users\gizquierdog\Documents\custom_report\custom_report_modificado.docx")
    print("✅ Documento generado correctamente con gráfica insertada.")
    doc_path = (r"C:\Users\gizquierdog\Documents\custom_report\custom_report_modificado.docx")
    return doc_path

